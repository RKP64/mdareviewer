import streamlit as st
import openai
from openai import AzureOpenAI
import json
import base64
import os
import pandas as pd
import tempfile
import html
import traceback
import re
import docx
import requests
from docx import Document
from docx.shared import Inches
import io
from dotenv import load_dotenv

# Imports for creating Word documents and parsing HTML
import markdown
from bs4 import BeautifulSoup

# Azure Search specific imports
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.models import VectorizedQuery

# Langchain specific imports for embeddings and chat models
from langchain_openai import AzureChatOpenAI as LangchainAzureChatOpenAI
from langchain_openai import AzureOpenAIEmbeddings
from langchain.schema import HumanMessage

# --- Page Configuration ---
# This must be the first Streamlit command. It sets the title, icon, and layout of the web page.
try:
    st.set_page_config(page_title="BIAL Multi-Agent Regulatory Platform", page_icon="✈️", layout="wide")
except Exception as e_config:
    print(f"CRITICAL ERROR during st.set_page_config: {e_config}")
    st.error(f"Error during st.set_page_config: {e_config}")
    st.stop()

# --- Helper Function to Check Credentials ---
# This function checks if a credential value is a placeholder, which helps prevent errors.
def check_creds(cred_value, placeholder_prefix="YOUR_"):
    if not cred_value: return True
    if isinstance(cred_value, str):
        if placeholder_prefix in cred_value.upper() or "ENTER_YOUR" in cred_value.upper() or (cred_value.startswith("<") and cred_value.endswith(">")):
            return True
    return False

# --- Main Application Logic ---
# The core function that runs the Streamlit application.
def main_app_logic():
    # --- Load Credentials from .env File ---
    # `load_dotenv()` reads the .env file and makes the variables available to `os.getenv()`.
    # This is the secure way to handle API keys without hardcoding them.
    load_dotenv()

    # Load all necessary credentials and configuration from environment variables.
    # If a variable isn't found, a default value is used.
    AZURE_SEARCH_ENDPOINT = os.getenv("AZURE_SEARCH_ENDPOINT")
    AZURE_SEARCH_API_KEY = os.getenv("AZURE_SEARCH_API_KEY")
    DEFAULT_AZURE_SEARCH_INDEX_NAME = os.getenv("DEFAULT_AZURE_SEARCH_INDEX_NAME", "mdmreviewer2")
    DEFAULT_VECTOR_FIELD_NAME = os.getenv("DEFAULT_VECTOR_FIELD_NAME", "contentVector")
    DEFAULT_SEMANTIC_CONFIG_NAME = os.getenv("DEFAULT_SEMANTIC_CONFIG_NAME", "azureml-default")

    AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
    # UPDATED: The API version is now set to a preview version compatible with o3-mini.
    AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION", "2024-03-01-preview")
    AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
    
    # Deployment for conversational tasks (default)
    AZURE_OPENAI_DEPLOYMENT_ID = os.getenv("AZURE_OPENAI_DEPLOYMENT_ID", "gpt-4o-mini")
    # Deployment specifically for long-form MDA reports (can be a model with a larger context window)
    AZURE_OPENAI_MDA_DEPLOYMENT_ID = os.getenv("AZURE_OPENAI_MDA_DEPLOYMENT_ID", "o3-mini")
    
    AZURE_OPENAI_PLANNING_DEPLOYMENT_ID = os.getenv("AZURE_OPENAI_PLANNING_DEPLOYMENT_ID", "gpt-4o-mini")
    AZURE_OPENAI_EMBEDDING_DEPLOYMENT_ID = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT_ID", "text-embedding-3-large")

    BING_SEARCH_API_KEY = os.getenv("BING_SEARCH_API_KEY")
    BING_SEARCH_ENDPOINT = os.getenv("BING_SEARCH_ENDPOINT", "https://api.bing.microsoft.com/v7.0/search")
    SERPAPI_API_KEY = os.getenv("SERPAPI_API_KEY")

    # --- Initialize API Clients ---
    # Configure the global OpenAI settings for Azure.
    openai.api_type = "azure"
    openai.api_base = AZURE_OPENAI_ENDPOINT
    openai.api_version = AZURE_OPENAI_API_VERSION
    openai.api_key = AZURE_OPENAI_API_KEY

    # Initialize clients for different AI models: embeddings, planning, and synthesis (final answer).
    search_query_embeddings_model, planning_openai_client, synthesis_openai_client = None, None, None
    try:
        # Initialize the model that turns text into vectors (embeddings) for searching.
        if not any(check_creds(c) for c in [AZURE_OPENAI_API_KEY, AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_EMBEDDING_DEPLOYMENT_ID, AZURE_OPENAI_API_VERSION]):
            search_query_embeddings_model = AzureOpenAIEmbeddings(azure_deployment=AZURE_OPENAI_EMBEDDING_DEPLOYMENT_ID, azure_endpoint=AZURE_OPENAI_ENDPOINT, api_key=AZURE_OPENAI_API_KEY, api_version=AZURE_OPENAI_API_VERSION, chunk_size=1)
    except Exception as e: st.sidebar.error(f"Error initializing Embeddings Model: {e}")
    try:
        # Initialize the clients for the main language models.
        if not any(check_creds(c) for c in [AZURE_OPENAI_API_KEY, AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_API_VERSION]):
            if not check_creds(AZURE_OPENAI_PLANNING_DEPLOYMENT_ID): planning_openai_client = AzureOpenAI(api_key=AZURE_OPENAI_API_KEY, azure_endpoint=AZURE_OPENAI_ENDPOINT, api_version=AZURE_OPENAI_API_VERSION)
            if not check_creds(AZURE_OPENAI_DEPLOYMENT_ID): synthesis_openai_client = AzureOpenAI(api_key=AZURE_OPENAI_API_KEY, azure_endpoint=AZURE_OPENAI_ENDPOINT, api_version=AZURE_OPENAI_API_VERSION)
    except Exception as e: st.sidebar.error(f"Error initializing OpenAI Clients: {e}")

    # --- Tool Functions (Web Search, Document Search, etc.) ---
    
    def query_bing_web_search(query: str, count: int = 5) -> str:
        # Function to perform a web search using the Bing Search API.
        if check_creds(BING_SEARCH_API_KEY) or check_creds(BING_SEARCH_ENDPOINT): return "Error: Bing Search API credentials are not configured."
        headers = {"Ocp-Apim-Subscription-Key": BING_SEARCH_API_KEY}
        params = {"q": query, "count": count, "textDecorations": True, "textFormat": "HTML"}
        try:
            response = requests.get(BING_SEARCH_ENDPOINT, headers=headers, params=params, timeout=10)
            response.raise_for_status()
            search_results = response.json()
            snippets = [f"Title: {res['name']}\nURL: {res['url']}\nSnippet: {res['snippet']}\n---" for res in search_results.get("webPages", {}).get("value", [])]
            return "\n".join(snippets) if snippets else "No Bing web search results found."
        except Exception as e: return f"Error during Bing web search: {e}"

    def query_serpapi(query: str, count: int = 5) -> str:
        # Function to perform a web search using the SerpApi service.
        if check_creds(SERPAPI_API_KEY):
            return "Error: SerpApi API key is not configured."
        params = {
            "q": query,
            "api_key": SERPAPI_API_KEY,
            "num": count,
            "engine": "google"
        }
        try:
            response = requests.get("https://serpapi.com/search.json", params=params, timeout=15)
            response.raise_for_status()
            search_results = response.json()
            organic_results = search_results.get("organic_results", [])
            snippets = [f"Title: {res.get('title', 'N/A')}\nURL: {res.get('link', 'N/A')}\nSnippet: {res.get('snippet', 'N/A')}\n---" for res in organic_results]
            return "\n".join(snippets) if snippets else "No web search results found via SerpApi."
        except Exception as e:
            return f"Error during SerpApi search: {e}"

    def get_query_vector(text_to_embed):
        # Converts a text query into a numerical vector using the embeddings model.
        if not search_query_embeddings_model:
            st.toast("Search Query Embedding model not ready.", icon="⚠️")
            return None
        try: return search_query_embeddings_model.embed_query(text_to_embed)
        except Exception as e: st.error(f"Error generating query vector: {e}"); return None

    @st.cache_data(ttl=3600)
    def get_indexes():
        # Retrieves the list of available search indexes from Azure Search. Caches the result for 1 hour.
        indexes = [DEFAULT_AZURE_SEARCH_INDEX_NAME] if DEFAULT_AZURE_SEARCH_INDEX_NAME and not check_creds(DEFAULT_AZURE_SEARCH_INDEX_NAME) else []
        if check_creds(AZURE_SEARCH_API_KEY) or check_creds(AZURE_SEARCH_ENDPOINT): return list(set(indexes))
        try:
            search_creds = AzureKeyCredential(AZURE_SEARCH_API_KEY)
            index_client = SearchIndexClient(AZURE_SEARCH_ENDPOINT, search_creds)
            indexes.extend([index.name for index in index_client.list_indexes()])
        except Exception as e: st.sidebar.error(f"Cannot retrieve Azure Search indexes: {e}", icon="🚨")
        return list(set(indexes))

    def query_azure_search(query_text, index_name, k=5, use_hybrid_semantic_search=True, vector_field_name=DEFAULT_VECTOR_FIELD_NAME, semantic_config_name=DEFAULT_SEMANTIC_CONFIG_NAME):
        # Performs a search on the specified Azure Search index.
        context, references_data = "", []
        if any(check_creds(c) for c in [AZURE_SEARCH_API_KEY, AZURE_SEARCH_ENDPOINT, index_name]): return "Error: Azure Search credentials or index name are placeholders.", []
        try:
            search_client = SearchClient(AZURE_SEARCH_ENDPOINT, index_name, AzureKeyCredential(AZURE_SEARCH_API_KEY))
            select_fields = ["content", "filepath", "url", "title"]
            search_kwargs = {"search_text": query_text if query_text and query_text.strip() else "*", "top": k, "include_total_count": True, "select": ",".join(select_fields)}
            if use_hybrid_semantic_search:
                if not check_creds(vector_field_name) and (query_vector := get_query_vector(query_text)):
                    search_kwargs["vector_queries"] = [VectorizedQuery(vector=query_vector, k_nearest_neighbors=k, fields=vector_field_name)]
                if not check_creds(semantic_config_name):
                    search_kwargs.update({"query_type": "semantic", "semantic_configuration_name": semantic_config_name, "query_caption": "extractive", "query_answer": "extractive"})
            
            results = search_client.search(**search_kwargs)
            if results.get_count() == 0: return "", []
            
            processed_references = {}
            for idx, doc in enumerate(results):
                doc_content = doc.get("content", "")
                context += doc_content + "\n\n"
                display_name = doc.get("title") or (os.path.basename(doc.get("filepath", "")) if doc.get("filepath") else f"Source {idx+1}")
                ref_key = doc.get("url") or display_name
                if ref_key not in processed_references:
                    processed_references[ref_key] = {
                        "filename_or_title": display_name, 
                        "url": doc.get("url"), 
                        "score": doc.get("@search.score"), 
                        "reranker_score": doc.get("@search.reranker_score"),
                        "content": doc_content
                    }
            references_data = list(processed_references.values())
        except Exception as e: return f"Error accessing search index '{index_name}': {e}", []
        return context.strip(), references_data

    def get_query_plan_from_llm(user_question, client_for_planning):
        # Uses an AI model to break down a complex user question into a series of simpler search queries.
        if not client_for_planning or check_creds(AZURE_OPENAI_PLANNING_DEPLOYMENT_ID): return "Error: Planning LLM not configured.", None
        planning_prompt = f"""You are a query planning assistant specializing in breaking down complex questions about **AERA regulatory documents, often concerning tariff orders, consultation papers, control periods, and specific financial data (like CAPEX, Opex, Traffic) for airport operators such as DIAL, MIAL, BIAL, HIAL.**
Your primary task is to take a user's complex question related to these topics and break it down into a series of 1 to 20 simple, self-contained search queries that can be individually executed against a document index. Each search query should aim to find a specific piece of information (e.g., a specific figure, a justification, a comparison point) needed to answer the overall complex question.
If the user's question is already simple and can be answered with a single search, return just that single query in the list.
If the question is very complex and might require more distinct search steps, formulate the most critical 1 to 20 search queries, focusing on distinct pieces of information.
Return your response ONLY as a JSON list of strings, where each string is a search query.
User's complex question: {user_question}
Your JSON list of search queries:"""
        try:
            response = client_for_planning.chat.completions.create(model=AZURE_OPENAI_PLANNING_DEPLOYMENT_ID, messages=[{"role": "user", "content": planning_prompt}], max_tokens=16000)
            plan_str = response.choices[0].message.content
            if match := re.search(r'\[.*\]', plan_str, re.DOTALL):
                query_plan = json.loads(match.group(0))
                if isinstance(query_plan, list) and all(isinstance(q, str) for q in query_plan): return None, query_plan
            return None, [user_question]
        except Exception as e: return f"Error getting query plan from LLM: {e}", None

    # --- CORE RAG (Retrieval-Augmented Generation) FUNCTION ---
    def generate_answer_from_search(user_question, index_name, use_hybrid_semantic, vector_field, semantic_config, max_tokens_param, client_for_synthesis, show_details=True, system_prompt_override=None, word_count_target=None, synthesis_deployment_id_override=None, document_context=""):
        if not client_for_synthesis: return "Error: Synthesis LLM not configured.", []
        
        synthesis_deployment = synthesis_deployment_id_override or AZURE_OPENAI_DEPLOYMENT_ID

        if show_details: st.write("⚙️ Generating query plan...")
        plan_error, query_plan = get_query_plan_from_llm(user_question, planning_openai_client)
        
        query_plan = query_plan or [user_question]
        if show_details:
            if plan_error: st.error(plan_error)
            st.write(f"📝 **Execution Plan:**")
            for i, q_step in enumerate(query_plan): st.write(f"   Step {i+1}: {q_step}")

        combined_context_for_llm = ""
        all_retrieved_details = []
        for i, sub_query in enumerate(query_plan):
            if show_details: st.write(f"🔍 Executing Step {i+1}: Searching for '{sub_query}'...")
            context_for_step, retrieved_details_for_step = query_azure_search(sub_query, index_name, use_hybrid_semantic_search=use_hybrid_semantic, vector_field_name=vector_field, semantic_config_name=semantic_config)
            if context_for_step and not context_for_step.startswith("Error"):
                combined_context_for_llm += f"\n\n--- Context for sub-query: '{sub_query}' ---\n" + context_for_step
                all_retrieved_details.extend(retrieved_details_for_step)
        
        if show_details: st.session_state.last_retrieved_chunks_search = all_retrieved_details

        # Combine the uploaded document context with the retrieved search context.
        full_context = ""
        if document_context:
            full_context += f"UPLOADED DOCUMENT CONTEXT:\n---------------------\n{document_context}\n---------------------\n\n"
        if combined_context_for_llm.strip():
            full_context += f"RETRIEVED KNOWLEDGE BASE CONTEXT:\n---------------------\n{combined_context_for_llm.strip()}\n---------------------\n\n"

        if not full_context: return "No relevant information found from any source.", []

        unique_sources_list = list({ (item.get('url') or item.get('filename_or_title')): item for item in all_retrieved_details }.values())
        formatted_refs_str = "\n".join([f"[doc{i+1}] {html.escape(item['filename_or_title'])}" for i, item in enumerate(unique_sources_list)])

        if show_details: st.write("💡 Synthesizing final answer...")
        
        # --- UPDATED PROMPT ---
        # Added a specific instruction to identify source tables.
        final_instructions = """**CRITICAL INSTRUCTIONS FOR YOUR RESPONSE:**
1.  **Answer Solely from Context:** Base your answer *exclusively* on the information found in the `AGGREGATED CONTEXT` section above.
2.  **Cite Everything:** You **MUST** provide an inline citation `[docN]` for every single piece of information, claim, or data point you use from the context. The number `N` must correspond to the source in the `IDENTIFIED CONTEXT SOURCES` list.
3.  **Format and Cite Tables:** When presenting data in a table, ensure it is valid HTML. Crucially, when extracting data from a source table, you must cite the source document and identify the original table by its title or number if possible.
4.  **Be Comprehensive:** Synthesize information from multiple sources to provide a full answer.
5.  **If Unsure, State It:** If the context does not contain the answer, state that explicitly. Do not use outside knowledge.
"""
        
        synthesis_prompt = (
            f"You are an AI assistant for Multiyear Tariff Submission for AERA.\n\n"
            f"USER QUESTION: {user_question}\n\n"
            f"IDENTIFIED CONTEXT SOURCES:\n---------------------\n{formatted_refs_str}\n---------------------\n\n"
            f"AGGREGATED CONTEXT:\n---------------------\n{full_context}\n---------------------\n\n"
            f"{final_instructions}\n\n"
        )
        
        if word_count_target:
            synthesis_prompt += f"Generate a response that is approximately **{word_count_target} words** long.\n\n"

        synthesis_prompt += "YOUR COMPREHENSIVE, CITED ANSWER:\n"
        
        try:
            # --- FIX: Use the correct parameter name based on the selected model ---
            completion_params = {
                "model": synthesis_deployment,
                "messages": [{"role": "user", "content": synthesis_prompt}],
                
            }
            if synthesis_deployment == "o3-mini":
                completion_params['max_completion_tokens'] = max_tokens_param
            else:
                completion_params['max_tokens'] = max_tokens_param

            response = client_for_synthesis.chat.completions.create(**completion_params)
            return response.choices[0].message.content, unique_sources_list
        except Exception as e: 
            return f"Error generating synthesized answer: {e}", []

    # --- Other Helper Functions (Report Refinement, Document Parsing) ---
    def refine_and_regenerate_report(original_report: str, new_info: str, client_for_synthesis) -> str:
        if not client_for_synthesis: return "Error: Synthesis LLM client not initialized."
        refinement_prompt = f"""You are a report writing expert. Your task is to seamlessly integrate a new piece of information into an existing report. Do not simply append the new information. Instead, find the most relevant section in the 'ORIGINAL REPORT' and intelligently merge the 'NEW INFORMATION' into it. Rewrite paragraphs as needed to ensure the final report is coherent, clean, and well-integrated. Return ONLY the full, updated report text.
        **ORIGINAL REPORT:**
        ---
        {original_report}
        ---
        **NEW INFORMATION TO INTEGRATE:**
        ---
        {new_info}
        ---
        **FULL, REFINED, AND INTEGRATED REPORT:**"""
        with st.spinner("✨ Refining report with new information..."):
            try:
                response = client_for_synthesis.chat.completions.create(model=AZURE_OPENAI_DEPLOYMENT_ID, messages=[{"role": "user", "content": refinement_prompt}], max_tokens=st.session_state.conv_agent_max_tokens)
                st.toast("Report successfully refined!", icon="✅")
                return response.choices[0].message.content
            except Exception as e:
                st.error(f"Error during report refinement: {e}")
                return original_report

    def extract_text_from_docx(file):
        try: return "\n".join([para.text for para in docx.Document(file).paragraphs])
        except Exception as e: st.error(f"Error reading Word document: {e}"); return None

    def parse_html_to_docx(soup, document):
        """Recursively parses BeautifulSoup elements and adds them to a docx document."""
        for element in soup.find_all(True, recursive=False):
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                try:
                    level = int(element.name[1])
                    document.add_heading(element.get_text(strip=True), level=level)
                except (ValueError, IndexError):
                    document.add_heading(element.get_text(strip=True), level=2)
            elif element.name == 'p':
                document.add_paragraph(element.get_text(strip=True))
            elif element.name == 'table':
                try:
                    rows = element.find_all('tr')
                    if not rows: continue
                    
                    headers = [th.get_text(strip=True) for th in rows[0].find_all(['th', 'td'])]
                    if not headers: continue
                    
                    table = document.add_table(rows=1, cols=len(headers))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = header
                    
                    for row in rows[1:]:
                        row_cells_data = [td.get_text(strip=True) for td in row.find_all('td')]
                        if len(row_cells_data) == len(headers):
                            row_cells = table.add_row().cells
                            for i, cell_text in enumerate(row_cells_data):
                                row_cells[i].text = cell_text
                except Exception as e:
                    document.add_paragraph(f"(Error parsing table: {e})")
            elif element.name in ['ul', 'ol']:
                style = 'List Bullet' if element.name == 'ul' else 'List Number'
                for li in element.find_all('li', recursive=False):
                    p = document.add_paragraph(style=style)
                    p.add_run(li.get_text(strip=True, separator=' ').split('\n')[0])
            elif element.name in ['div', 'section', 'article']:
                 parse_html_to_docx(element, document)

    def create_word_document(markdown_text):
        """
        Creates a Word document from markdown text, correctly parsing HTML tags.
        """
        try:
            document = Document()
            styles = document.styles
            try:
                if 'List Bullet' not in styles:
                    styles.add_style('List Bullet', 1).base_style = styles['List Paragraph']
                if 'List Number' not in styles:
                    styles.add_style('List Number', 1).base_style = styles['List Paragraph']
            except Exception as e:
                print(f"Could not add default list styles: {e}")

            html_content = markdown.markdown(markdown_text, extensions=['markdown.extensions.tables'])
            soup = BeautifulSoup(html_content, 'html.parser')
            
            parse_html_to_docx(soup, document)

            file_stream = io.BytesIO()
            document.save(file_stream)
            file_stream.seek(0)
            return file_stream
        except Exception as e:
            st.error(f"Failed to create Word document: {e}")
            document = Document()
            document.add_paragraph("Error creating document.")
            file_stream = io.BytesIO()
            document.save(file_stream)
            file_stream.seek(0)
            return file_stream

    @st.cache_data
    def get_base64_image(image_path: str):
        try:
            with open(image_path, "rb") as f: return base64.b64encode(f.read()).decode("utf-8")
        except Exception: return None

    def display_answer_with_citations(answer_text, sources):
        if not sources or not isinstance(answer_text, str) or not re.search(r'\[doc\d+\]', answer_text):
            st.markdown(answer_text, unsafe_allow_html=True)
            return

        parts = re.split(r'(\[doc\d+\])', answer_text)
        
        for part in parts:
            part = part.strip()
            if not part:
                continue

            if match := re.match(r'\[doc(\d+)\]', part):
                doc_num = int(match.group(1))
                source_index = doc_num - 1
                
                if 0 <= source_index < len(sources):
                    source = sources[source_index]
                    with st.expander(f"[{doc_num}]"):
                        st.markdown(f"**Source:** {source.get('filename_or_title', 'N/A')}")
                        st.caption(source.get('content', 'Content not available.'))
            else:
                st.markdown(part, unsafe_allow_html=True)


    # --- UI & State Management ---
    st.markdown(""" <style>
    :root {
        --primary-background: #33263C; /* Dominant dark purple/brown background */
        --card-background: #2A2A2A; /* Dark gray of the card/table header */
        --text-color-light: #E0E0E0; /* Light gray for most text */
        --text-color-dark: #A0A0A0; /* Slightly darker gray for some text or secondary information */
        --header-text-color: #FFFFFF; /* White for the header "Deploy" and column titles */
        --button-green: #28A745; /* Green color of the "Generate" button */
        --sidebar-dark: #1E1E1E; /* Very dark almost black for the sidebar area */
    }

    /* General Styles */
    body {
        color: var(--text-color-light);
    }
    .stApp {
        background-color: var(--primary-background);
    }

    /* Main content area (response box, text area, etc.) */
    .response-box, .stTextArea, .stTextInput, .st-expander {
        background-color: var(--card-background) !important;
        border: 1px solid #4A4A4A !important; /* A slightly lighter border for definition */
        color: var(--text-color-light) !important;
        border-radius: 8px;
    }
    .st-expander header {
        color: var(--text-color-light) !important;
    }

    .response-box {
        padding: 20px;
        margin-bottom: 20px;
        white-space: pre-wrap;
        word-wrap: break-word;
    }
    
    /* History Entry in Sidebar */
    .history-entry {
        border: 1px solid #4A4A4A;
        padding: 15px;
        margin-bottom: 15px;
        border-radius: 8px;
        background-color: var(--card-background);
        white-space: pre-wrap;
        word-wrap: break-word;
    }

    /* Titles and Headers */
    .title-bar h1, h1, h2, h3, h4, p {
        color: var(--header-text-color);
    }
    
    /* --- CSS FOR ALL BUTTONS TO BE GREEN --- */
    .stButton>button {
        border: none;
        background-color: var(--button-green);
        color: var(--header-text-color);
        border-radius: 8px;
        padding: 12px 24px;
        font-weight: bold;
        transition: all 0.3s ease-in-out;
    }

    .stButton>button:hover {
        background-color: #218838; /* A darker green for hover */
        box-shadow: 0 0 15px rgba(40, 167, 69, 0.5);
    }
    
    .stButton>button:active {
        background-color: #1E7E34 !important; /* An even darker green for active/click */
    }

    div[data-testid="column"] .stButton>button {
        height: 100%;
        width: 100%;
        text-align: left !important;
        padding: 15px;
        font-weight: normal;
    }

    div[data-testid="column"] .stButton>button:hover {
        transform: translateY(-3px);
    }
    
    /* --- END OF BUTTON CSS --- */

    .st-emotion-cache-1r6slb0 { /* Sidebar */
        background-color: var(--sidebar-dark);
    }
    .st-emotion-cache-16txtl3 { /* Main content area */
        background-color: var(--primary-background);
    }
    
    </style>
    
    
    """, unsafe_allow_html=True)
    
    default_session_state = {
        "conversation_history": [], "last_retrieved_chunks_search": [], "question_text": "",
        "conv_agent_temp": 0.5, "conv_agent_max_tokens": 16000, "conv_agent_selected_index": DEFAULT_AZURE_SEARCH_INDEX_NAME,
        "conv_agent_use_hybrid": True, "conv_agent_vector_field": DEFAULT_VECTOR_FIELD_NAME,
        "conv_agent_semantic_config": DEFAULT_SEMANTIC_CONFIG_NAME, "app_mode": "Conversational Agent",
        "mda_analysis_type": "MDA Manpower Analysis", "mda_report_content": None, "mda_report_sources": [],
        "mda_chat_history": [], "web_search_engine": "Bing", "mda_word_count": 4000,
        "mda_synthesis_model": "gpt-4o-mini", # Default model for MDA reports
        "custom_analysis_prompt": "" # State for the custom prompt text area
    }
    for key, value in default_session_state.items():
        if key not in st.session_state: st.session_state[key] = value

    with st.sidebar:
        st.sidebar.title("Choose App")
        st.session_state.app_mode = st.sidebar.radio("Select the application to use:", ("Conversational Agent", "MDA Reviewer"), key="app_mode_selector")
        st.markdown("---")
        
        with st.expander("⚙️ Settings", expanded=True):
            st.header("Agent Settings")
            st.session_state.conv_agent_temp = st.slider("Temperature", 0.0, 1.0, st.session_state.conv_agent_temp, 0.1)
            st.session_state.conv_agent_max_tokens = st.slider("Output", 100, 16000, st.session_state.conv_agent_max_tokens, 50)
            
            st.session_state.web_search_engine = st.radio("Web Search Engine", ("Bing", "SerpApi"))

            st.subheader("Regulatory database")
            available_indexes = get_indexes()
            st.session_state.conv_agent_selected_index = st.selectbox("Select Index", available_indexes, index=0 if not st.session_state.conv_agent_selected_index in available_indexes else available_indexes.index(st.session_state.conv_agent_selected_index))
            st.session_state.conv_agent_use_hybrid = st.checkbox("Enable Multi-Agent Flow", value=st.session_state.conv_agent_use_hybrid)
            if st.session_state.conv_agent_use_hybrid:
                st.session_state.conv_agent_vector_field = st.text_input("Vector Field Name", value=st.session_state.conv_agent_vector_field)
                st.session_state.conv_agent_semantic_config = st.text_input("Semantic Configuration", value=st.session_state.conv_agent_semantic_config)
            
            # --- Model selection for MDA Reviewer ---
            if st.session_state.app_mode == "MDA Reviewer":
                st.markdown("---")
                st.subheader("MDA Reviewer Settings")
                st.session_state.mda_synthesis_model = st.radio(
                    "Select Model for Report Generation",
                    ["gpt-4o-mini", "o3-mini"],
                    key="mda_model_selector",
                    help="Choose the model for generating the detailed MDA report. 'o3-mini' may be better for very long documents."
                )

        if st.session_state.app_mode == "Conversational Agent":
            st.markdown("---")
            st.subheader("📜 Chat History")
            if st.session_state.conversation_history:
                for entry in reversed(st.session_state.conversation_history):
                    st.markdown(f"**Q:** {entry['question']}", unsafe_allow_html=True)
                    st.markdown(f"<div class='history-entry'>{entry['answer']}</div>", unsafe_allow_html=True)
                    st.markdown("---", unsafe_allow_html=True)

                if st.button("Clear History"): 
                    st.session_state.conversation_history = []
                    st.rerun()

    col_header1, col_header2 = st.columns([1, 10])
    with col_header1:
        if main_logo_base64 := get_base64_image("bial_logo.png"): st.image(f"data:image/png;base64,{main_logo_base64}", width=400)
    with col_header2: st.markdown('<div class="title-bar"><h1>BIAL Regulatory Assistant</h1></div>', unsafe_allow_html=True)
    st.markdown("<hr style='margin-top: 0; margin-bottom:1em;'>", unsafe_allow_html=True)

    if st.session_state.app_mode == "Conversational Agent":
        st.subheader(f"💬 Conversational Agent (Index: {st.session_state.conv_agent_selected_index or 'Not Selected'})")
        
        for entry in st.session_state.conversation_history:
            with st.chat_message("user"):
                st.markdown(entry["question"])
            with st.chat_message("assistant"):
                with st.container():
                     display_answer_with_citations(entry["answer"], entry.get("sources", []))
        
        if not st.session_state.conversation_history:
            st.markdown("---")
            st.subheader("Ask a question")
            predefined_questions = [
                "Calculate and compare the YoY change of employee expenses of DIAL and MIAL for the fourth control period",
                "What is the YoY change of employee expenses submitted by MIAL for the fourth control period and the rationale for the growth rates",
                "Compare the manpower expense per total passenger traffic submitted by DIAL and MIAL respectively for fourth control period."
            ]
            cols = st.columns(len(predefined_questions))
            for i, q in enumerate(predefined_questions):
                if cols[i].button(q, key=f"predef_q_{i}"):
                    st.session_state.conversation_history.append({"role": "user", "question": q, "answer": "Thinking...", "sources": []})
                    st.rerun()

        user_query = st.chat_input("Ask a question, e.g., 'Compare manpower expenses for DIAL and MIAL.'")
        if user_query:
            st.session_state.conversation_history.append({"role": "user", "question": user_query, "answer": "Thinking...", "sources": []})
            st.rerun()

    if st.session_state.app_mode == "Conversational Agent" and st.session_state.conversation_history and st.session_state.conversation_history[-1]["answer"] == "Thinking...":
        latest_entry = st.session_state.conversation_history[-1]
        user_query = latest_entry["question"]
        
        with st.spinner("Thinking..."):
            answer_text, sources = generate_answer_from_search(
                user_question=user_query,
                index_name=st.session_state.conv_agent_selected_index,
                use_hybrid_semantic=st.session_state.conv_agent_use_hybrid,
                vector_field=st.session_state.conv_agent_vector_field,
                semantic_config=st.session_state.conv_agent_semantic_config,
            
                max_tokens_param=st.session_state.conv_agent_max_tokens,
                client_for_synthesis=synthesis_openai_client,
                show_details=False
            )
            latest_entry["answer"] = answer_text
            latest_entry["sources"] = sources
            st.rerun()

    elif st.session_state.app_mode == "MDA Reviewer":
        st.subheader("📄 Review and validation of MDA")
        st.info("Analyze and validate BIAL's regulatory submissions by benchmarking against peer airports and historical data.", icon="ℹ️")

        analysis_prompts_config = {
            "MDA Manpower fourth control period Analysis": {
                "Comprehensive Manpower Analysis": f"""
                1- projected personnel costs, including closing manpower, cost per employee, and total personnel costs for each fiscal year by BIAL for 4th control period along with justification provided by BIAL.
                2- caluclate year on year Manpower Expenses growth Submitted by DIAL for fourth control period in DIAL fourth control period consultation Paper along with justification provide by DIAL and Authority
                3- calculate Year on Year growth of employee cost submitted by MIAL for fourth control period for fourth control period in MIAL Fourth control consultation Paper along with justification provided by MIAL and Authority .
                4- Caluclate and compare  growth rate of year on year personnel expenditure of BIAL with the growth rates projected by DIAL and MIAL for fourth control period 
                5- suggest how the rationale or justification provided by BIAL in the MDA_personnel document for manpower expenditure for fourth control period can be enhanced by checking into rationale provided by MIAL and DIAL for their fourth control period Submission and the what authority has said on DIAL and MIAL Submission. For every suggestion made, give specific reason why the suggestion was made by you using relevant references from DIAL and MIAL tariff orders or consultation papers. 
                      
                """,
            },


             "MDA Manpower third control period  Analysis": {
                "Comprehensive Manpower  Analysis": f"""
                1- Actual manpower expenditure for BIAL and variance from authority approved manpower expenditure for the third control period along with justification provided by BIAL and Authority 
                2- Actual manpower expenditure for DIAL and variance from authority approved manpower expenditure for the third control period alomg with the justification provided by DIAl and authority.
                3-Actual manpower expenditure for MIAL and variance from authority approved manpower expenditure for the third control period along with justification provided by MIAL and Authority
                4-suggest how the rationale or justification provided by BIAL in the MDA_personnel document for manpower expenditure for third control period can be enhanced by checking into rationale provided by MIAL and DIAL for their control period Submission and the what authority has said on DIAL and MIAL Submission. For every suggestion made, give specific reason why the suggestion was made by you using relevant references from DIAL and MIAL third control period  tariff orders or consultation paper .
                """,
             },

            "Utility Analysis for Electricity": {
                "Comprehensive Electricity Analysis for third control period": f"""
                1- year on year actual power consumption actual recoveries from power consumption by BIAL in third control period in utlities_mda _document along with justification provided by BIAL in the utlities_mda_document. 
                2- From the fourth control period consultation paper, extract the year-on-year electricity consumption and the corresponding recovery from the electricity sub-concessionaire submitted by DIAL for the true-up of the third control period, including the justifications provided by both DIAL and the Authority for these figures
                3- From the MIAL fourth control period consultation paper, provide the year-on-year electricity cost, gross consumption, and recoveries submitted by MIAL for the true-up of the third control period, along with MIAL's justification for the variance in actual electricity cost from the authority's approved figures and the authority's own examination and rationale on the matter.
                4- Conduct a comparative analysis of electricity metrics from the third control period true-up for DIAL, MIAL, and BIAL. Specifically, compare DIAL's electricity consumption, MIAL's electricity cost, and BIAL's actual power cost consumption  in utlities_mda document . Based on this data, formulate a detailed rationale that explains any significant trends or performance differences between the airports, providing clear reasons for your conclusions.
                """,
                
            }, 


             "Utility Analysis for water": {
                "Comprehensive Water Analysis for third control period": f"""
                1- From the Utlities_cost_MDA document for the third control period true-up, extract the year-on-year actual cost, consumption, and recoveries for portable and raw water, along with BIAL's complete justification for the water expense and the variance from the authority's approved figures. 
                2- Year on Year  water gross charge submitted by DIAL for true up of third control period in the DIAL fourth control period consultation paper. 6- Year on Year growth of water consumption submitted by DIAL for third control period in the DIAL fourth control period consultation paper. 7- Year on Year actual recoveries from sub- concessionaire submitted by DIAL for third control period in the DIAL fourth control period consultation paper. 8- Justification for actual  gross water charge  in third control period in the DIAL fourth control period consultation paper provided by DIAL and the variance with authority approved figures. 9- Examination and rationale provided by authority on actual water gross charge and consumption submitted by DIAL for third control period in the DIAL fourth control period consultation paper.
                3-From the MIAL fourth control period consultation paper for the true-up of the third control period, extract the year-on-year water expense, consumption, and recoveries, along with the justification provided by MIAL for the variance from approved figures, and the Authority's final examination and rationale on the submission.
                4-Develop a set of actionable recommendations to enhance the justification for BIAL's third control period water expenditure in Utilities_cost_MDA Document. For each recommendation, provide a specific reason, citing relevant best practices from the justifications and rationale provided by DIAL and MIAL in their respective fourth consultation paper.
                """,
                
            }, 
            
            
            # Empty dictionaries for category placeholders
            "R&M Analysis for third control Period": {
               "comprehensive R&M Analysis for third control period":f"""
               1-From BIAL's MDA_O&M document for the third control period, extract the year-on-year actual growth of repairs and maintenance (R&M) expenditure, the annual R&M cost as a percentage of the regulated asset base, and BIAL's complete justification for the expense, including the variance from the authority's approved figures
               2-From the fourth control period consultation paper or tariff order, extract DIAL's year-on-year growth of actual repairs and maintenance (R&M) expenditure for the third control period true-up, the year-wise R&M expense as a percentage of its regulated asset base, and the justification provided for the variance from the authority's approved figures.
               3-From the fourth control period consultation paper or tariff order for the true-up of the third control period, extract MIAL's year-on-year growth of actual repairs and maintenance (R&M) expenditure, the R&M expense as a percentage of its regulated asset base, the justification provided by MIAL for the variance from approved figures, and the Authority's final examination and rationale on the matter
               4-Develop a set of actionable recommendations to enhance the justification for BIAL's third control period repairs and maintenance expenditure, using a comparative analysis of the rationale provided by DIAL and MIAL in their respective tariff orders or consultation papers to support each suggestion with specific reasons and references.
               """,
            },
               "R&M Analysis for fourth control Period":{
                "comprehensiive R&M Aanlysis for fourth control period":f"""
                1-From BIAL's MDA_O&M document, extract the projected year-on-year growth of repairs and maintenance (R&M) expenditure for the fourth control period, the projected R&M cost as a percentage of the regulated asset base, and the complete justification for these projections.
                2-From the fourth control period consultation paper or tariff order, extract DIAL's projected year-on-year growth of repairs and maintenance (R&M) expenditure, the projected R&M cost as a percentage of the regulated asset base, the justification provided by DIAL for these projections, and the Authority's final examination and rationale on the submission.
                3-From the fourth control period consultation paper or tariff order, extract MIAL's projected year-on-year growth of repairs and maintenance (R&M) expenditure, the projected R&M cost as a percentage of the regulated asset base, the justification provided by MIAL for these projections, and the Authority's final examination and rationale on the submission.
                4-Develop a set of actionable recommendations to enhance the justification for BIAL's fourth control period repairs and maintenance expenditure projections. For each recommendation, provide a specific reason, citing relevant best practices from the justifications and rationale provided by DIAL, MIAL, and the Authority in their respective fourth control period tariff orders or consultation papers.

              """,              
               },

            }
        
        
        st.session_state.mda_analysis_type = st.selectbox("Select Analysis Category:", list(analysis_prompts_config.keys()))
        
        # --- UPDATED LOGIC FOR HYBRID APPROACH ---
        # Get the predefined prompts for the selected category.
        analysis_options_dict = analysis_prompts_config.get(st.session_state.mda_analysis_type, {})
        if isinstance(analysis_options_dict, dict):
            specific_analysis_options = list(analysis_options_dict.keys())
        else:
            specific_analysis_options = []
        
        # Add the custom prompt option to the list.
        custom_prompt_option = "Custom Analysis Prompt..."
        specific_analysis_options.append(custom_prompt_option)
        
        selected_specific_analysis_title = st.selectbox("Select Specific Analysis:", specific_analysis_options)

        # Show a text area only if the custom option is selected.
        custom_prompt_text = ""
        if selected_specific_analysis_title == custom_prompt_option:
            custom_prompt_text = st.text_area(
                "Enter your custom multi-step analysis prompt below:",
                height=250,
                placeholder="Enter your instructions here. Each step must start with a number and a hyphen (e.g., '1- ...').\n\nExample:\n1- Analyze the YoY growth of personnel cost for BIAL.\n2- Compare BIAL's growth to DIAL's growth.",
                key="custom_analysis_prompt"
            )

        st.session_state.mda_word_count = st.number_input("Target Word Count for Report", min_value=100, max_value=6000, value=st.session_state.get("mda_word_count", 6000), step=500)
        uploaded_word_file = st.file_uploader("Upload a Word Document (.docx)", type="docx")

        if st.button("Generate Specific Analysis Report", type="primary", disabled=(not uploaded_word_file or (selected_specific_analysis_title == custom_prompt_option and not custom_prompt_text))):
            st.session_state.mda_report_content, st.session_state.mda_chat_history = None, []
            with st.spinner("Processing document and generating analysis..."):
                if (extracted_text := extract_text_from_docx(uploaded_word_file)):
                    
                    # Determine which prompt to use (predefined or custom).
                    prompt_template = ""
                    if selected_specific_analysis_title == custom_prompt_option:
                        prompt_template = custom_prompt_text
                    else:
                        prompt_template = analysis_prompts_config.get(st.session_state.mda_analysis_type, {}).get(selected_specific_analysis_title, "")

                    if prompt_template:
                        # --- BATCH-WISE, STEP-BY-STEP WORKFLOW ---
                        report_parts = [f"## Analysis Report for: *{html.escape(uploaded_word_file.name)}*"]
                        all_sources = []
                        
                        # FIX: The regex now correctly handles multi-line prompts.
                        analysis_steps = [s.strip() for s in re.findall(r'^\s*\d+-\s*(.*)', prompt_template, re.M)]
                        
                        # DYNAMIC MODEL SELECTION
                        selected_model_name = st.session_state.mda_synthesis_model
                        mda_deployment_id = AZURE_OPENAI_MDA_DEPLOYMENT_ID if selected_model_name == "o3-mini" else AZURE_OPENAI_DEPLOYMENT_ID

                        with st.status(f"Running multi-step analysis for '{selected_specific_analysis_title}'...", expanded=True) as status:
                            for i, step_text in enumerate(analysis_steps):
                                status.write(f"**Processing Step {i+1}/{len(analysis_steps)}:** {step_text}")
                                
                                # The user_question for each step is just the step text.
                                # The document_context is passed separately.
                                step_answer, step_sources = generate_answer_from_search(
                                    user_question=step_text,
                                    document_context=extracted_text,
                                    index_name=st.session_state.conv_agent_selected_index,
                                    use_hybrid_semantic=True,
                                    vector_field=st.session_state.conv_agent_vector_field,
                                    semantic_config=st.session_state.conv_agent_semantic_config,
                                
                                    max_tokens_param=st.session_state.conv_agent_max_tokens,
                                    client_for_synthesis=synthesis_openai_client,
                                    show_details=False,
                                    synthesis_deployment_id_override=mda_deployment_id
                                )
                                
                                report_parts.append(f"### Step {i+1}: {step_text}\n{step_answer}")
                                all_sources.extend(step_sources)
                            
                            # Final Refinement Step.
                            status.write("Combining and refining final report...")
                            combined_report_text = "\n\n---\n\n".join(report_parts)
                            
                            refinement_prompt = f"""You are an expert report writer. Based on the following detailed sections, please synthesize them into a single, cohesive, and well-structured final report. Ensure a logical flow, remove any redundancies, and format the output professionally. The final report should be approximately {st.session_state.mda_word_count} words.
                            
                            DETAILED SECTIONS TO SYNTHESIZE:
                            ---
                            {combined_report_text}
                            ---
                            
                            FINAL, COHESIVE REPORT:"""

                            # The final refinement does not need to search the knowledge base again.
                            final_report, _ = generate_answer_from_search(
                                user_question=refinement_prompt,
                                index_name=st.session_state.conv_agent_selected_index, 
                                use_hybrid_semantic=False,
                                vector_field="",
                                semantic_config="",
                            
                                max_tokens_param=st.session_state.conv_agent_max_tokens,
                                client_for_synthesis=synthesis_openai_client,
                                show_details=False,
                                synthesis_deployment_id_override=mda_deployment_id
                            )

                            st.session_state.mda_report_content = final_report
                            st.session_state.mda_report_sources = list({v['filename_or_title']:v for v in all_sources}.values())
                            status.update(label="Analysis complete!", state="complete", expanded=False)

                        st.success("Analysis report generated successfully!")

        if st.session_state.mda_report_content:
            st.markdown("---")
            
            word_file = create_word_document(st.session_state.mda_report_content)
            st.download_button(
                label="📥 Download Report as Word",
                data=word_file,
                file_name=f"BIAL_Analysis_Report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            display_answer_with_citations(st.session_state.mda_report_content, st.session_state.mda_report_sources)
            
            st.markdown("---")
            st.subheader("💬 Follow-up Chat")
            for i, message in enumerate(st.session_state.mda_chat_history):
                with st.chat_message(message["role"]):
                    display_answer_with_citations(message["content"], message.get("sources", []))
                    if message["role"] == "assistant":
                        if st.button("✨ Integrate & Refine Report", key=f"refine_{i}"):
                            st.session_state.mda_report_content = refine_and_regenerate_report(st.session_state.mda_report_content, message["content"], synthesis_openai_client)
                            st.session_state.mda_chat_history = []
                            st.rerun()
            
            if prompt := st.chat_input("Ask to elaborate or find new information..."):
                st.session_state.mda_chat_history.append({"role": "user", "content": prompt, "sources": []})
                st.rerun()

    if st.session_state.app_mode == "MDA Reviewer" and st.session_state.mda_chat_history and st.session_state.mda_chat_history[-1]["role"] == "user":
        
        user_message = st.session_state.mda_chat_history[-1]
        prompt = user_message["content"]

        with st.spinner("Agent is thinking..."):
            answer = None
            sources = []
            prompt_lower = prompt.lower()
            
            web_search_keywords = ["bing search", "search bing", "google search", "serper", "search google", "latest", "current", "web search", "internet"]
            is_web_search = any(w in prompt_lower for w in web_search_keywords)
            
            if is_web_search:
                st.info(f"⚡ Performing Web Search ({st.session_state.web_search_engine})...")
                
                search_tool = query_bing_web_search
                if "SerpApi" in st.session_state.web_search_engine:
                    search_tool = query_serpapi
                
                context_from_search = search_tool(prompt)
                
                if context_from_search.startswith("Error:"):
                    answer = context_from_search
                else:
                    web_search_prompt = f"""You are an expert AI assistant. Based on the following SEARCH RESULTS, provide a concise and informative answer to the user's question.
                    USER QUESTION: "{prompt}"
                    SEARCH RESULTS:\n---\n{context_from_search}\n---\nYOUR DETAILED, FORMATTED ANSWER:"""
                    
                    response = synthesis_openai_client.chat.completions.create(model=AZURE_OPENAI_DEPLOYMENT_ID, messages=[{"role": "user", "content": web_search_prompt}]).choices[0].message.content
                    answer = response
            else:
                st.info(f"🧠 Searching Internal Documents...")
                answer, sources = generate_answer_from_search(
                    user_question=prompt,
                    index_name=st.session_state.conv_agent_selected_index,
                    use_hybrid_semantic=st.session_state.conv_agent_use_hybrid,
                    vector_field=st.session_state.conv_agent_vector_field,
                    semantic_config=st.session_state.conv_agent_semantic_config,
                
                    max_tokens_param=st.session_state.conv_agent_max_tokens,
                    client_for_synthesis=synthesis_openai_client,
                    show_details=False
                )
            
            st.session_state.mda_chat_history.append({"role": "assistant", "content": answer, "sources": sources})
            st.rerun()

if __name__ == '__main__':
    try:
        main_app_logic()
    except Exception as e:
        st.error(f"An critical unexpected error occurred: {e}")
        traceback.print_exc()







